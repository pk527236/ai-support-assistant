__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')

import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
import json

# FIX for Python 3.8 + ChromaDB telemetry issue
import chromadb
from chromadb.config import Settings

def load_text_file(file_path):
    """Load plain text files"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def load_docx_file(file_path):
    """Load Word documents (.docx)"""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except ImportError:
        print("⚠️ python-docx not installed. Install with: pip install python-docx")
        raise
    except Exception as e:
        print(f"❌ Error reading .docx file: {e}")
        raise

def load_pdf_file(file_path):
    """Load PDF files"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except ImportError:
        print("⚠️ PyPDF2 not installed. Install with: pip install PyPDF2")
        raise
    except Exception as e:
        print(f"❌ Error reading PDF file: {e}")
        raise

def load_json_file(file_path):
    """Load JSON files (for Teams chat history and Zendesk articles)"""
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Format Zendesk articles
    if isinstance(data, list) and len(data) > 0 and 'title' in data[0]:
        text = ""
        for article in data:
            title = article.get('title', 'Unknown')
            content = article.get('content', '')
            url = article.get('url', '')
            text += f"ARTICLE: {title}\n"
            if url:
                text += f"SOURCE: {url}\n"
            text += f"{'='*80}\n{content}\n\n"
        return text
    
    # Format Teams chat history
    elif isinstance(data, list):
        text = ""
        for msg in data:
            sender = msg.get('sender', 'Unknown')
            content = msg.get('content', '')
            text += f"{sender}: {content}\n\n"
        return text
    
    return str(data)

def load_documents_from_directory(directory):
    """Load all documents from a directory"""
    documents = []
    
    # Skip certain files
    skip_files = ['.env', 'README.md', '.gitignore', '.DS_Store']
    skip_dirs = ['__pycache__', '.git', 'chroma_db', 'zendesk_articles']
    
    for root, dirs, files in os.walk(directory):
        # Skip certain directories
        dirs[:] = [d for d in dirs if d not in skip_dirs]
        
        for file in files:
            if file in skip_files:
                continue
            
            # Skip hidden files
            if file.startswith('.'):
                continue
                
            file_path = os.path.join(root, file)
            
            try:
                # Text files
                if file.endswith('.txt'):
                    content = load_text_file(file_path)
                    documents.append(Document(
                        page_content=content, 
                        metadata={"source": file, "type": "text"}
                    ))
                    print(f"✅ Loaded: {file} ({len(content)} chars)")
                
                # JSON files
                elif file.endswith('.json'):
                    content = load_json_file(file_path)
                    doc_type = "zendesk" if "zendesk" in file.lower() else "json"
                    documents.append(Document(
                        page_content=content, 
                        metadata={"source": file, "type": doc_type}
                    ))
                    print(f"✅ Loaded: {file} ({len(content)} chars)")
                
                # Markdown files
                elif file.endswith('.md'):
                    content = load_text_file(file_path)
                    documents.append(Document(
                        page_content=content, 
                        metadata={"source": file, "type": "markdown"}
                    ))
                    print(f"✅ Loaded: {file} ({len(content)} chars)")
                
                # Word documents
                elif file.endswith('.docx'):
                    content = load_docx_file(file_path)
                    documents.append(Document(
                        page_content=content, 
                        metadata={"source": file, "type": "word"}
                    ))
                    print(f"✅ Loaded: {file} ({len(content)} chars)")
                
                # PDF files
                elif file.endswith('.pdf'):
                    content = load_pdf_file(file_path)
                    documents.append(Document(
                        page_content=content, 
                        metadata={"source": file, "type": "pdf"}
                    ))
                    print(f"✅ Loaded: {file} ({len(content)} chars)")
            
            except Exception as e:
                print(f"❌ Error loading {file}: {e}")
    
    return documents

def get_embeddings():
    """Get embeddings function with fallback options"""
    
    # Option 1: Try HuggingFace embeddings (Recommended - Free & Offline)
    try:
        from langchain_community.embeddings import HuggingFaceEmbeddings
        print("🔧 Using HuggingFace embeddings...")
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'}
        )
        print("✅ HuggingFace embeddings loaded")
        return embeddings
    except Exception as e:
        print(f"⚠️ HuggingFace embeddings failed: {e}")
    
    # Option 2: Try Ollama embeddings
    try:
        from langchain_community.embeddings import OllamaEmbeddings
        print("🔧 Using Ollama embeddings...")
        embeddings = OllamaEmbeddings(
            model="llama3.2",
            base_url="http://localhost:11434"
        )
        print("✅ Ollama embeddings loaded")
        return embeddings
    except Exception as e:
        print(f"⚠️ Ollama embeddings failed: {e}")
        raise Exception("Could not load any embedding model. Please install sentence-transformers: pip install sentence-transformers")

def ingest_data():
    """Main function to ingest and process data"""
    
    # Check if data directory exists
    data_dir = "./data"
    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        print(f"📁 Created {data_dir} directory")
        print("⚠️ Please add your training data files and run again")
        print("\nTo scrape Zendesk content:")
        print("  1. Run: python zendesk_scraper.py")
        print("  2. Then run this script again")
        return
    
    # Load documents
    print("📚 Loading documents from ./data directory...")
    documents = load_documents_from_directory(data_dir)
    
    if not documents:
        print("⚠️ No documents found in ./data directory")
        print("\nExpected file types: .txt, .json, .md, .docx, .pdf")
        print("\nTo add Zendesk content:")
        print("  1. Run: python zendesk_scraper.py")
        print("  2. Then run this script again")
        return
    
    print(f"\n✅ Loaded {len(documents)} documents")
    
    # Show document types
    doc_types = {}
    for doc in documents:
        dtype = doc.metadata.get('type', 'unknown')
        doc_types[dtype] = doc_types.get(dtype, 0) + 1
    
    print("\n📊 Document breakdown:")
    for dtype, count in doc_types.items():
        print(f"  - {dtype}: {count} files")
    
    # Split documents into chunks
    print("\n✂️ Splitting documents into chunks...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
    )
    
    splits = text_splitter.split_documents(documents)
    print(f"✅ Created {len(splits)} chunks")
    
    # Create embeddings with fallback
    print("\n🧠 Creating embeddings (this may take a few minutes)...")
    print("⏳ First run will download the embedding model (~90MB)")
    embeddings = get_embeddings()
    
    # Create and persist vector store
    persist_directory = "./chroma_db"
    
    # Delete existing vector store if it exists
    if os.path.exists(persist_directory):
        import shutil
        print(f"🗑️ Removing old vector store...")
        shutil.rmtree(persist_directory)
    
    # Configure ChromaDB settings to disable telemetry
    chroma_settings = Settings(
        anonymized_telemetry=False,
        allow_reset=True
    )
    
    print("💾 Creating vector store...")
    
    # Create persistent client with settings
    client = chromadb.PersistentClient(
        path=persist_directory,
        settings=chroma_settings
    )
    
    vector_store = Chroma.from_documents(
        documents=splits,
        embedding=embeddings,
        persist_directory=persist_directory,
        client=client
    )
    
    print(f"\n✅ Vector store created and saved to {persist_directory}")
    print("="*80)
    print("✅ TRAINING COMPLETE!")
    print("="*80)
    print(f"📊 Statistics:")
    print(f"  - Documents processed: {len(documents)}")
    print(f"  - Chunks created: {len(splits)}")
    print(f"  - Vector store location: {persist_directory}")
    print(f"\n🚀 Next step: Run 'python app.py' to start your AI assistant!")

if __name__ == "__main__":
    print("="*80)
    print("AI Support Bot - Knowledge Base Ingestion")
    print("="*80)
    print("\nSupported file types: .txt, .json, .md, .docx, .pdf")
    print()
    ingest_data()